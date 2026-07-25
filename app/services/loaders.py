"""
文件加载器: 把上传的原始文件转成 Segment 列表
  Segment = {"type": "text"|"image", "content": str}
    - text:  content 是纯文本
    - image: content 是 data-url (data:image/png;base64,...), chunk 表存 "[IMAGE:xxx]" 占位

支持的 loader:
  text       .txt/.log/.json/.csv/...   纯文本(utf-8)
  markdown   .md/.markdown              去 Markdown 标记
  pdf_fast   .pdf                       PyMuPDF 快速抽文本(不抽图片)
  pdf_deep   .pdf                       PyMuPDF 深度: 文本 + 图片(转base64 data-url), 可走多模态向量
  docx       .docx                      python-docx 抽段落+表格
"""
from __future__ import annotations
from abc import ABC, abstractmethod
import os
import re
import io
import base64
from dataclasses import dataclass


@dataclass
class Segment:
    type: str           # "text" | "image"
    content: str        # 文本 or data-url
    page: int = 0       # 页码(用于溯源)


class BaseLoader(ABC):
    extensions: tuple[str, ...] = ()
    name: str = "base"

    def load(self, raw: bytes) -> list[Segment]:
        """读取原始字节, 返回 Segment 列表"""
        return [Segment(type="text", content=raw.decode("utf-8", errors="ignore"))]


class TextLoader(BaseLoader):
    extensions = (".txt", ".log", ".json", ".csv", ".xml", ".yaml", ".yml", ".py", ".js", ".ts", ".html", ".css")
    name = "text"

    def load(self, raw: bytes) -> list[Segment]:
        text = raw.decode("utf-8", errors="ignore")
        return [Segment(type="text", content=text)]


class MarkdownLoader(BaseLoader):
    extensions = (".md", ".markdown")
    name = "markdown"

    def load(self, raw: bytes) -> list[Segment]:
        text = raw.decode("utf-8", errors="ignore")
        text = re.sub(r"```[\s\S]*?```", " ", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"!\[[^\]]*\]\([^)]*\)", " ", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
        text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
        text = re.sub(r"~~([^~]+)~~", r"\1", text)
        text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
        return [Segment(type="text", content=text.strip())]


class PDFFastLoader(BaseLoader):
    """快速 PDF: 只用 PyMuPDF 抽文字, 不解析图片。"""
    extensions = (".pdf",)
    name = "pdf_fast"

    def load(self, raw: bytes) -> list[Segment]:
        import fitz
        doc = fitz.open(stream=raw, filetype="pdf")
        try:
            segments: list[Segment] = []
            for i, page in enumerate(doc):
                text = page.get_text("text") or ""
                text = text.strip()
                if text:
                    segments.append(Segment(type="text", content=text, page=i + 1))
            return segments
        finally:
            doc.close()


class PDFDeepLoader(BaseLoader):
    """深度 PDF: PyMuPDF 抽文字 + 图片。图片转 PNG base64, 作为 image Segment。
       后续会用 embed_multimodal 对图片段单独向量化,文本段走文本 embedding。"""
    extensions = (".pdf",)
    name = "pdf_deep"

    def load(self, raw: bytes) -> list[Segment]:
        import fitz
        doc = fitz.open(stream=raw, filetype="pdf")
        try:
            segments: list[Segment] = []
            for i, page in enumerate(doc):
                # 1) 文本
                text = (page.get_text("text") or "").strip()
                if text:
                    segments.append(Segment(type="text", content=text, page=i + 1))
                # 2) 图片(每页的 image list)
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha >= 4:  # CMYK -> RGB
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        png_bytes = pix.tobytes("png")
                        b64 = base64.b64encode(png_bytes).decode("ascii")
                        data_url = f"data:image/png;base64,{b64}"
                        segments.append(Segment(type="image", content=data_url, page=i + 1))
                    except Exception:
                        pass
            return segments
        finally:
            doc.close()


class DocxLoader(BaseLoader):
    extensions = (".docx",)
    name = "docx"

    def load(self, raw: bytes) -> list[Segment]:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(raw))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return [Segment(type="text", content="\n".join(parts).strip())]


class UnstructuredLoader(BaseLoader):
    """Unstructured: 多格式通用解析(PDF/DOCX/HTML/PPT/XLSX/EML/...)
       依赖 unstructured 库(pip install unstructured[pdf,docx]等)。
       速度较慢、依赖重,适合格式混杂或前面专用 loader 失败时兜底。
       默认不被 auto 选中,需用户手动选 loader=unstructured。
    """
    # 不绑定扩展名; 用户显式选才用
    extensions = ()
    name = "unstructured"

    def load(self, raw: bytes) -> list[Segment]:
        try:
            from unstructured.partition.auto import partition
        except ImportError:
            raise RuntimeError("缺少 unstructured 依赖,请 pip install 'unstructured[pdf,docx]'")
        # partition 接收 file 字节流
        elements = partition(file=io.BytesIO(raw))
        # 把所有 element 的 text 合并(去重空行)
        texts = []
        for el in elements:
            t = (el.text or "").strip()
            if t:
                texts.append(t)
        return [Segment(type="text", content="\n".join(texts))]


# 兼容旧 pdf loader 名 (默认快速)
class PDFLoader(PDFFastLoader):
    name = "pdf"


# ---------- 注册表 ----------
LOADERS: dict[str, BaseLoader] = {}


def register(loader: BaseLoader, aliases: tuple[str, ...] = ()):
    LOADERS[loader.name] = loader
    for a in aliases:
        LOADERS[a] = loader


register(TextLoader(), aliases=("plain", "txt"))
register(MarkdownLoader(), aliases=("md",))
register(PDFFastLoader())
register(PDFDeepLoader())
register(PDFLoader())        # 'pdf' 别名 -> 快速
register(DocxLoader())
register(UnstructuredLoader())


def get_loader(name: str) -> BaseLoader:
    if name in LOADERS:
        return LOADERS[name]
    lower = name.lower()
    for loader in set(LOADERS.values()):
        if lower in loader.extensions:
            return loader
    raise ValueError(f"未知 loader: {name}, 可选: {list_available()}")


def list_available() -> list[dict]:
    """返回 [{key, label, desc}], 前端按此展示下拉"""
    return [
        {"key": "auto",        "label": "自动识别",     "desc": "按扩展名自动选择加载器"},
        {"key": "text",        "label": "纯文本",       "desc": ".txt/.json/.csv 等"},
        {"key": "markdown",    "label": "Markdown",     "desc": "去 Markdown 标记保留正文"},
        {"key": "pdf_fast",    "label": "PDF 快速",     "desc": "PyMuPDF 抽文字,速度快"},
        {"key": "pdf_deep",    "label": "PDF 深度",     "desc": "抽文字+图片,多模态向量化"},
        {"key": "docx",        "label": "Word (docx)",  "desc": "段落+表格"},
        {"key": "unstructured","label": "Unstructured","desc": "多格式兜底(PDF/DOCX/HTML/PPT),慢但通用"},
    ]


def detect_by_filename(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return "pdf_fast"     # 默认快速;深度需要用户明确选
    for loader in set(LOADERS.values()):
        if ext in loader.extensions and loader.name != "pdf":
            return loader.name
    return "text"


def segments_to_plain_text(segments: list[Segment]) -> str:
    """把 text 段拼成大文本(用于非多模态流程的旧函数兼容)"""
    return "\n\n".join(s.content for s in segments if s.type == "text")
